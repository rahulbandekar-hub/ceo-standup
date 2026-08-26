#!/usr/bin/env python3
"""
Weekly Primeharvest leadership digest — fully local replacement for the cloud routine
that failed silently twice (completed with no error, but never wrote to Drive).

Does real research via Claude's server-side web_search tool and writes the result
straight to ~/Desktop/Prime Harvest/ as a .docx. No cloud session, no Drive
dependency — every step runs and can be verified locally.

Cost-trimmed 2026-08-05 (target: well under $0.50/week): dropped the "What's Changed
Since Last Week" section and the previous-digest lookup that fed it (fewer input
tokens, one less thing to break), cut personas to 1-2 tight bullets each instead of
2-3, cut web search from 20 to 8 uses, moved to the dynamic-filtering search tool
(web_search_20260318, which filters results server-side before they hit the context
window instead of dumping raw pages in) with response_inclusion "excluded" to drop
already-consumed search content from the reply, and cut max_tokens from 16000 to
6000 to match the now-shorter output. Estimated cost per run: ~$0.10-0.15
(previously ~$1-3 per run — real-time pricing, 20 uncapped basic searches, up to
16000 output tokens across a much longer 9-persona + CEO's-call + What's-Changed
structure).
"""
import json
import os
import re
import sys
from datetime import datetime, timezone
from pathlib import Path

import anthropic
from docx import Document
from docx.shared import Pt

HERE = Path(__file__).parent
ENV_PATH = HERE / ".env"
DESKTOP_DIR = Path.home() / "Desktop" / "Prime Harvest"
LEDGER_PATH = Path.home() / "api-spend-watchdog" / "ledger.jsonl"

def load_api_key():
    for line in ENV_PATH.read_text().splitlines():
        if line.startswith("ANTHROPIC_API_KEY="):
            return line.split("=", 1)[1].strip()
    raise RuntimeError("ANTHROPIC_API_KEY not found in .env")

def log_spend(usage):
    """Local spend tracking for ~/api-spend-watchdog — this account stays personal,
    which has no Admin API / Cost Report API access, so the watchdog can't ask
    Anthropic what was spent. Log raw token usage here instead; the watchdog does
    the $ math itself."""
    try:
        LEDGER_PATH.parent.mkdir(parents=True, exist_ok=True)
        entry = {
            "ts": datetime.now(timezone.utc).isoformat(),
            "source": "primeharvest-digest",
            "model": "claude-sonnet-5",
            "service_tier": "standard",
            "input_tokens": usage.input_tokens,
            "output_tokens": usage.output_tokens,
            "cache_read_input_tokens": getattr(usage, "cache_read_input_tokens", 0) or 0,
            "cache_creation_input_tokens": getattr(usage, "cache_creation_input_tokens", 0) or 0,
            "web_search_requests": getattr(getattr(usage, "server_tool_use", None), "web_search_requests", 0) or 0,
        }
        with open(LEDGER_PATH, "a") as f:
            f.write(json.dumps(entry) + "\n")
    except Exception as e:
        print(f"  log_spend: failed to log spend (non-fatal): {e}", file=sys.stderr)

PROMPT_TEMPLATE = """You are producing this week's "leadership meeting minutes" for PRIMEHARVEST, a
fictitious CPG company invented by Rahul (an MIT Sloan MBA graduate with ~10 years of CPG commercial
strategy experience at General Mills, now CEO of this fictitious company). Rahul uses this to stay
sharp on the real US Consumer Products (CPG) and Retail industry while interviewing for senior
CPG/retail roles.

COMPANY PROFILE (for framing only — do not invent news about this fictitious company; use it only as
the lens through which real industry news is discussed):
Primeharvest is a ~$18B multi-category packaged foods company, publicly traded, Midwest HQ. Portfolio
spans breakfast/cereal, sweet & salty snacks, baking mixes, frozen meals, and a smaller pet food
division. #1 or #2 share in most of its categories. Category captain or co-captain at Walmart, Kroger,
and Target in 2-3 categories. Currently under pressure: private label share is climbing, snacks are
seeing GLP-1-linked volume softness, and two years of price increases have started eroding volume and
consumer trust. The standing leadership tension: Finance wants margin discipline; Sales/Marketing want
investment to rebuild share and trust, without another blunt price hike.

GOAL: Every real news item discussed must be grounded in something you actually found via web search
from the past 7 days — cite the source/outlet and approximate date for every point. Do not invent
trends, statistics, or company actions. This is not fiction for its own sake; Primeharvest and its
leadership team are simply the lens for discussing what's genuinely happening in the real US CPG/retail
industry this week.

Use web search to find real news from the past 7 days covering: consumer trends (spending shifts,
category growth/decline, private label, value-seeking behavior, GLP-1 impact on food, health/wellness
reformulation), major CPG companies (PepsiCo, Coca-Cola, General Mills, Kraft Heinz, Mondelez, Nestle,
Unilever, P&G, Conagra, Kellanova, Hershey, Mars, StarKist) and major retailers (Walmart, Kroger,
Target, Costco, Amazon) — competitive moves, pricing/promotion strategy, market share shifts, retail
media growth, earnings commentary; M&A activity, divestitures, leadership changes; supply chain
disruptions; regulatory/policy news (labeling, FDA/USDA, tariffs, litigation), ESG/sustainability moves;
agentic AI / data infrastructure adoption in CPG and retail; R&D/innovation news (reformulation, new
product launches, functional/health-forward products). Search efficiently — 8-10 well-targeted searches
covering the widest useful spread of the topics above, not one search per topic.

Write the digest in this exact structure, using Markdown (## for section headers, ### for each
leader's name, - for bullets):

## Primeharvest Weekly Intel — {today}

## Leadership Meeting Minutes

For each of these 9 people, write 1-2 tight bullets (one line each, no sub-bullets), EACH grounded in a
specific real news item (name the company/outlet and approximate date), written in their voice/lens on
what it means for Primeharvest's growth, margin, share, or consumer trust:
- Marcus Webb, SVP Sales — retailer relationships, shelf space, customer goodwill
- Priya Chen, CFO — margin protection, ROI discipline, cost pressure
- Jordan Ellis, CMO (Brand & Marketing) — brand equity/trust AND in-market campaign/media activity
- Diego Ruiz, VP Supply Chain — input costs, capacity, disruption risk
- Dr. Elena Kovacs, VP Consumer Insights — behavioral/sentiment shifts, what consumers actually want now
- Aisha Patel, VP Revenue Growth Management — pricing, promo, mix, elasticity
- Tom Reyes, VP Category Management — shelf/assortment, category captain dynamics, competitive share
- Raj Malhotra, Chief Data & Technology Officer — IT infrastructure, agentic AI adoption, data ecosystem
- Dr. Naomi Osei, Chief R&D & Innovation Officer — reformulation, new products, functional/health innovation

## External Affairs (this week)
Only include Victor Adeyemi, Chief External Affairs & Strategy Officer (regulatory/policy,
ESG/sustainability, M&A/divestitures) if there is genuinely relevant real news this week — otherwise
omit this entire section.

## CEO's Call
Written in Rahul's voice as CEO: a short, sharp synthesis (3-4 sentences) reacting to the week and
naming the 1-2 things he's actually prioritizing, given the standing tension between margin discipline
and rebuilding share/trust.

## Sources
List every source cited above with outlet name, headline/topic, and date.

TONE: Real internal meeting notes — concise, functional, specific, distinct voice per person. No fluff.
If you can't find enough real news for a function this week, say so briefly rather than padding.

IMPORTANT — paraphrase, don't lift: write every point in the leader's own voice and words. Do not copy
sentences directly from articles or search results. Any direct quote (e.g. an executive's public remark)
must be short (under 15 words), in quotation marks, and clearly attributed — never quote a journalist's
or analyst's descriptive sentence as if reporting a fact in your own words.
"""

def build_prompt():
    today = datetime.now().strftime("%B %d, %Y")
    return PROMPT_TEMPLATE.format(today=today)

TITLE_MARKER = "## Primeharvest Weekly Intel"

def call_claude(prompt, api_key):
    client = anthropic.Anthropic(api_key=api_key)
    resp = client.messages.create(
        model="claude-sonnet-5",
        max_tokens=6000,
        tools=[{
            "type": "web_search_20260318",
            "name": "web_search",
            "max_uses": 8,
            "response_inclusion": "excluded",
        }],
        messages=[{"role": "user", "content": prompt}],
    )
    log_spend(resp.usage)
    text_parts = [block.text for block in resp.content if block.type == "text"]
    full_text = "\n".join(text_parts)

    # The model sometimes narrates in-progress ("I'll research...") in an early text
    # block before it starts using tools — strip everything before the actual digest
    # title so that preamble never ends up in the saved document.
    marker_idx = full_text.find(TITLE_MARKER)
    if marker_idx > 0:
        full_text = full_text[marker_idx:]

    if resp.stop_reason == "max_tokens":
        print(
            "WARNING: response was truncated by max_tokens — digest is likely incomplete. "
            "Consider raising max_tokens further.",
            file=sys.stderr,
        )
        full_text += "\n\n[NOTE: This digest was cut off mid-generation and is incomplete.]"

    return full_text

def markdown_to_docx(md_text, out_path):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for raw_line in md_text.splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=0)
        elif re.match(r"^[-*]\s+", line):
            content = re.sub(r"^[-*]\s+", "", line)
            content = re.sub(r"\*\*(.+?)\*\*", r"\1", content)  # strip bold markers, keep text
            doc.add_paragraph(content, style="List Bullet")
        else:
            content = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
            doc.add_paragraph(content)

    doc.save(str(out_path))

def main():
    api_key = load_api_key()
    if not api_key:
        print("ERROR: ANTHROPIC_API_KEY is empty in .env", file=sys.stderr)
        sys.exit(1)

    DESKTOP_DIR.mkdir(parents=True, exist_ok=True)

    print(f"[{datetime.now().isoformat()}] Building prompt...")
    prompt = build_prompt()

    print(f"[{datetime.now().isoformat()}] Calling Claude with web search (this takes a few minutes)...")
    digest_md = call_claude(prompt, api_key)

    if not digest_md.strip():
        print("ERROR: Empty response from Claude — aborting without writing a file", file=sys.stderr)
        sys.exit(1)

    if "## Sources" not in digest_md:
        print(
            "ERROR: Response has no Sources section — digest is incomplete or malformed. "
            "Not saving, so a broken file never silently replaces a good one.",
            file=sys.stderr,
        )
        sys.exit(1)

    date_str = datetime.now().strftime("%B %d %Y")
    out_path = DESKTOP_DIR / f"Primeharvest Weekly Intel - {date_str}.docx"
    markdown_to_docx(digest_md, out_path)

    print(f"[{datetime.now().isoformat()}] Saved: {out_path}")

if __name__ == "__main__":
    main()
